import re
import pypandoc
import pypandoc.pandoc_download
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def md_string_to_docx(
    md_content: str,
    output_file: str = "output.docx",
    title: str = None,
):
    """
    Convert Markdown text into a professional Word (.docx) file:
    - Removes sequences of '===' lines (3 or more).
    - Uses Calibri font (official style).
    - Keeps heading hierarchy (Heading 1, Heading 2, etc.).
    - Normalizes spacing for professional look.
    - Adds title at the top of the document.
    - No footer / date added.
    """

    # 🔹 Remove sequences of 3+ '='
    cleaned_content = re.sub(r"={3,}", "", md_content)

    # Ensure Pandoc installed
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        print("📦 Pandoc not found. Downloading...")
        pypandoc.pandoc_download.download_pandoc()

    # Convert Markdown → temporary DOCX
    temp_file = "temp_content.docx"
    pypandoc.convert_text(
        cleaned_content,
        to="docx",
        format="md",
        outputfile=temp_file,
        extra_args=["--standalone"],
    )

    # Open converted docx
    doc = Document(temp_file)

    # 🔹 Normalize formatting
    for para in doc.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15

        if para.style.name.startswith("Heading"):
            for run in para.runs:
                run.font.name = "Calibri"
        else:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(12)

    # 🔹 Wrap with title if provided
    if title:
        new_doc = Document()

        # Add title
        p = new_doc.add_paragraph(title)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(20)
        run.bold = True

        new_doc.add_paragraph("")  # spacing after title

        # Merge content
        for element in doc.element.body:
            new_doc.element.body.append(element)

        doc = new_doc  # replace with wrapper

    # Save final file
    doc.save(output_file)
    return output_file  # ✅ Return the saved file path

# """
#     md_string_to_docx(
#         md_content,
#         output_file="Functional_Specs.docx",
#         title="📄 Functional Specification Document - Logical Unit-1",
#     )




# --- for PDF Geneartion ---
import os
import markdown2
from xhtml2pdf import pisa

def md_string_to_pdf(md_content: str, output_file: str = "output.pdf", title: str = None, output_path: str = None):
    """
    Convert Markdown to styled PDF using markdown2 + xhtml2pdf
    """
    # 🔹 Remove sequences of 3+ '='
    md_content = re.sub(r"={3,} *", "", md_content)

    # Convert Markdown → HTML
    html_content = markdown2.markdown(md_content, extras=["tables", "fenced-code-blocks"])

    # Wrap with HTML/CSS
    html_template = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Calibri, Helvetica, Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.5;
                margin: 50px;
            }}
            /* ✅ First line (title) centered */
            .doc-title {{
                text-align: center;
                font-size: 20pt;
                font-weight: bold;
                margin-bottom: 20px;
                color: #003366; /* Dark Blue */
            }}
            /* ✅ Heading colors */
            h1 {{ font-size: 18pt; margin: 20px 0 10px 0; color: #003366; }}
            h2 {{ font-size: 16pt; margin: 18px 0 8px 0; color: #006600; }}
            h3 {{ font-size: 14pt; margin: 16px 0 6px 0; color: #993300; }}
            p {{ margin: 6px 0; }}
            ul {{ margin: 6px 0 6px 20px; }}
            li {{ margin: 4px 0; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 12px 0;
            }}
            th, td {{
                border: 1px solid #444;
                padding: 6px 10px;
                text-align: left;
            }}
            th {{
                background: #eee;
            }}
            code {{
                font-family: Consolas, monospace;
                background: #f4f4f4;
                padding: 2px 4px;
                border-radius: 3px;
            }}
            pre {{
                background: #f4f4f4;
                padding: 10px;
                border-radius: 5px;
                overflow-x: auto;
            }}
            blockquote {{
                border-left: 4px solid #ccc;
                margin: 10px 0;
                padding-left: 12px;
                color: #555;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        {f"<div class='doc-title'>{title}</div>" if title else ""}
        {html_content}
    </body>
    </html>
    """
    # Convert HTML → PDF
    output_file_path = output_path/output_file
    with open(output_file_path, "wb") as f:
        pisa_status = pisa.CreatePDF(html_template, dest=f)

    if pisa_status.err:
        raise Exception("Error during PDF generation")

    print(f"✅ PDF saved at {output_file_path}")
    return output_file_path








